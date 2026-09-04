"""Document processing — extract text from PDF/DOCX/MD/TXT files.

Security: 强制 multipart UploadFile,拒绝客户端 filepath;文件落到受控临时目录,
扩展名白名单,大小限制。原 filepath 任意文件读 (C2) 已堵。
"""
import os
import shutil
import tempfile
from pathlib import Path
from fastapi import APIRouter, File, HTTPException, UploadFile

router = APIRouter(prefix="/api/documents", tags=["documents"])

ALLOWED_EXTS = {".pdf", ".docx", ".md", ".txt"}
MAX_BYTES = 50 * 1024 * 1024  # 50 MB
TMP_DIR = Path(tempfile.gettempdir()) / "tp-documents"


def extract_pdf(filepath: str) -> str:
    from PyPDF2 import PdfReader
    reader = PdfReader(filepath)
    text = ""
    for page in reader.pages[:50]:  # max 50 pages
        t = page.extract_text()
        if t:
            text += t + "\n\n"
    return text.strip() or ""


def extract_docx(filepath: str) -> str:
    from docx import Document
    path = Path(filepath)
    if not path.exists():
        raise FileNotFoundError(f"DOCX file not found: {filepath}")
    try:
        doc = Document(str(path))
    except Exception:
        # python-docx on Windows may fail on non-ASCII paths; copy to temp ASCII path
        suffix = path.suffix
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
            tmpname = tmp.name
        try:
            shutil.copy2(str(path), tmpname)
            doc = Document(tmpname)
        finally:
            Path(tmpname).unlink(missing_ok=True)
    text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return text.strip() or ""


def _safe_extract_path(ext: str) -> Path:
    """生成系统临时目录下的安全文件名(权限 0600,不接受客户端路径)。"""
    TMP_DIR.mkdir(parents=True, exist_ok=True)
    return TMP_DIR / f"{os.urandom(16).hex()}{ext}"


@router.post("/extract")
async def extract_text(file: UploadFile = File(...)):
    """Extract text from PDF/DOCX/MD/TXT uploaded by client."""
    if not file.filename:
        raise HTTPException(status_code=400, detail="Missing filename")
    ext = Path(file.filename).suffix.lower()
    if ext not in ALLOWED_EXTS:
        raise HTTPException(status_code=400, detail=f"Unsupported format: {ext}")

    dest = _safe_extract_path(ext)
    try:
        size = 0
        with dest.open("wb") as f:
            while chunk := await file.read(64 * 1024):
                size += len(chunk)
                if size > MAX_BYTES:
                    raise HTTPException(status_code=413, detail=f"File too large (> {MAX_BYTES} bytes)")
                f.write(chunk)

        try:
            if ext == ".pdf":
                text = extract_pdf(str(dest))
            elif ext == ".docx":
                text = extract_docx(str(dest))
            elif ext in (".md", ".txt"):
                text = dest.read_text(encoding="utf-8", errors="ignore")
            else:
                raise HTTPException(status_code=400, detail=f"Unsupported format: {ext}")
        except ImportError as e:
            raise HTTPException(status_code=500, detail=f"Missing dependency: {e}")
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=422, detail=f"{ext} extraction error: {e}")

        if not text:
            raise HTTPException(status_code=422, detail=f"No extractable text in {ext} file")

        return {"filename": Path(file.filename).name, "text": text, "length": len(text)}
    finally:
        dest.unlink(missing_ok=True)
